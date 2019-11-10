from flask import Flask, render_template, request,flash, send_from_directory
from flask_wtf import Form
import os
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
app.secret_key = 'development key'

@app.route('/')
def form():
   return render_template('form.html')

@app.route('/pdf',methods = ['POST', 'GET'])
def pdf():
    if request.method == 'POST':
      result = request.form

      datetimeObject = datetime.strptime(result["Date"],'%Y-%m-%d')
      newformat = datetimeObject.strftime('%d-%m-%Y')
      
      s1 = "Dear Mr/Ms {},"
      s2 = "Hope you are well. I have decided to visit you on {} as per you request. This is just a heads-up."
      s3 = "Yours truely,"
      s4 = "Team JLT"
      
      pdf = FPDF()
      pdf.add_page()
      pdf.set_font("Arial", size=12)
      pdf.cell(150, 10, txt=s1.format(result["Name"]), ln=1, align="L")
      pdf.cell(0, 10, txt=s2.format(newformat), ln=1, align="C")
      pdf.cell(190, 10, txt=s3, ln=1, align="R")
      pdf.cell(190, 10, txt=s4, ln=1, align="R")
      pdf.output("Letter.pdf")

    try:
        return send_from_directory(os.path.dirname(os.path.abspath(__file__)),'Letter.pdf',as_attachment=True)
    except Exception as e:
        return str(e)

@app.route('/docx',methods = ['POST', 'GET'])
def docx():
    if request.method == 'POST':
      result = request.form

      datetimeObject = datetime.strptime(result["Date"],'%Y-%m-%d')
      newformat = datetimeObject.strftime('%d-%m-%Y')

      document = Document()

      document.add_heading('Letter',0)
      p1 = document.add_paragraph('Dear Mr/Ms ')
      p1.add_run(result["Name"]).bold = True
      p1.add_run(',')

      p2 = document.add_paragraph('Hope you are well. I have decided to visit you on ')
      p2.add_run(newformat).bold = True
      p2.add_run(' as per you request. This is just a heads-up.')

      document.add_paragraph('Yours truely,')
      document.add_paragraph('XXX')

      document.save('Letter.docx')

    try:
        return send_from_directory(os.path.dirname(os.path.abspath(__file__)),'Letter.docx',as_attachment=True)
    except Exception as e:
        return str(e)
      

if __name__ == '__main__':
   app.run(debug = True)

    